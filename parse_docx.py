from docx import Document
import json

# 读取 Word 文档
doc = Document('鑫正辉公司主营产品介绍.docx')

# 提取所有段落
content = []
for para in doc.paragraphs:
    if para.text.strip():
        content.append(para.text)

# 提取表格信息
tables = []
for table in doc.tables:
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        table_data.append(row_data)
    tables.append(table_data)

# 输出结果
print("=" * 80)
print("文档内容提取")
print("=" * 80)
print(f"\n共提取 {len(content)} 个段落\n")

for i, text in enumerate(content, 1):
    print(f"[{i}] {text}")
    print()

if tables:
    print("=" * 80)
    print(f"共提取 {len(tables)} 个表格\n")
    for i, table in enumerate(tables, 1):
        print(f"表格 {i}:")
        for row in table:
            print(" | ".join(row))
        print()

# 保存为 JSON
output = {
    'paragraphs': content,
    'tables': tables
}

with open('docx_content.json', 'w', encoding='utf-8') as f:
    json.dump(output, f, ensure_ascii=False, indent=2)

print("\n内容已保存到 docx_content.json")
